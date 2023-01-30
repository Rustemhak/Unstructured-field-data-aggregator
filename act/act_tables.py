import pandas as pd
import pdfplumber
from collections import OrderedDict
from pdf2docx import Converter
from docx.api import Document
from tabulate import tabulate


def clear_enter(text: str) -> str:
    """
    Очищает текст от символов переноса строки и соединяет слова, если присутсвует знак переноса
    :param text: сырой текст
    :return: обработанный и "чистый" текст
    """
    if text is None:
        return text
    # отдельно убирает знаки переноса
    text = text.replace('-\n', '')
    text = text.replace('\n', '')
    # text = text.replace('   ', '')
    # text = text.replace('  ','')
    # text = text.replace(' ', '')
    return text


def read_tables_act(page: pdfplumber.pdf.Page, table_indexes) -> list[pd.DataFrame]:
    """

    :return: возвращает таблицы  1. ОСНОВНЫЕ ДАННЫЕ О РЕМОНТЕ, 7.ПЕРФОРАЦИЯ, ОТКЛЮЧЕНИЕ ПЛАСТОВ, 16.РАСХОД МАТЕРИАЛОВ, ИСПОЛЬЗУЕМЫХ ПРИ РЕМОНТЕ
    """
    tables = page.extract_tables()
    tables = [tables[ind] for ind in table_indexes]
    tables_clean = []
    for table in tables:
        t = list(map(lambda x: list(map(lambda y: clear_enter(y), x)), table))
        tables_clean.append(t)
    df_list = []
    for table in tables_clean:
        df = pd.DataFrame(table)
        df_list.append(df)
    return df_list


# pdf = pdfplumber.open('МУН/Азнакаевскнефть/AKT_KRS_2836_АЗН.PDF')
def process_tales(path):
    pdf = pdfplumber.open(path)
    p0 = pdf.pages[0]
    table_settings = {
        # "layout" : True,
        # "use_text_flow" :True
        # "horizontal_ltr" : False
    }
    # tables = p0.extract_tables(table_settings)
    tables = p0.extract_tables(table_settings)
    table_indexes = [0, 4, 5]

    tables = [tables[ind] for ind in table_indexes]
    table0 = tables[0]
    table0 = list(map(lambda x: list(map(lambda y: clear_enter(y), x)), table0))
    # print(table0)
    features = table0[0]
    d_8_10 = OrderedDict.fromkeys([features[0], features[3], features[-1]])
    values = table0[2]
    d_8_10['Вид работы'] = [values[1]]
    d_8_10['Метод работы'] = [values[4]]
    d_8_10['Причина ремонта'] = [values[5]]
    for k, v in d_8_10.items():
        print(k, ':', v)

    # РАСХОД МАТЕРИАЛОВ, ИСПОЛЬЗУЕМЫХ ПРИ РЕМОНТЕ
    table2 = tables[2]
    table2 = list(map(lambda x: list(map(lambda y: clear_enter(y), x)), table2))
    df2 = pd.DataFrame(table2[1:], columns=table2[0])
    # print(df2)
    df2.to_excel('РАСХОД МАТЕРИАЛОВ, ИСПОЛЬЗУЕМЫХ ПРИ РЕМОНТЕ.xlsx')

    # ПЕРФОРАЦИЯ, ОТКЛЮЧЕНИЕ ПЛАСТОВ
    table1 = tables[1]
    table1 = list(map(lambda x: list(map(lambda y: clear_enter(y), x)), table1))
    df1 = pd.DataFrame(table1[1:], columns=table1[0])
    # print(df1)
    df1.to_excel('ПЕРФОРАЦИЯ, ОТКЛЮЧЕНИЕ ПЛАСТОВ.xlsx')

    df_8_10 = pd.DataFrame(data=d_8_10)
    return df_8_10, df1, df2


if __name__ == '__main__':
    # process_tales('AKT_KRS_2850_АЗН_пример.PDF')
    test_path_pdf = 'AKT_KRS_2850_АЗН_пример.PDF'

    result_path = 'test_result.docx'

    c = Converter(test_path_pdf)
    c.convert(result_path)
    c.close()
    document = Document(result_path)
    for tb_id in range(len(document.tables)):
        table = document.tables[tb_id]  # здесь нужно указать номер таблицы
        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (clear_enter(cell.text) for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
        print(data)

        df = pd.DataFrame(data)
        print(tabulate(df, headers='keys', tablefmt='psql'))

    # pdf = pdfplumber.open("table_test_pdf4.pdf")
    # p0 = pdf.pages[0]
    # tables = p0.extract_tables()
    # table0 = tables[0]
    # # table0 = list(map(lambda x: list(map(lambda y: clear_enter(y), x)), table0))
    # print(table0)
